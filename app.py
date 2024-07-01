from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import os
import docx
import pdfplumber
import spacy
from flask_cors import CORS
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

app = Flask(__name__)
CORS(app)
app.config['UPLOAD_FOLDER'] = 'uploads'

# Load SpaCy model
nlp = spacy.load('en_core_web_md')

def extract_text_from_pdf(pdf_path):
    text = ''
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ''
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    text += ' '.join(row) + ' '
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + ' '
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + ' '
    return text

def extract_keywords(text):
    doc = nlp(text)
    keywords = Counter()
    stop_words = set(nlp.Defaults.stop_words)

    for token in doc:
        if not token.is_stop and not token.is_punct and len(token.text) > 1:
            if token.pos_ in ['NOUN', 'ADJ', 'PROPN'] or token.ent_type_:
                lemma = token.lemma_.lower()
                if lemma not in stop_words and not lemma.isdigit():
                    keywords[lemma] += 1

    sorted_keywords = dict(sorted(keywords.items(), key=lambda item: item[1], reverse=True))
    return sorted_keywords

def calculate_score(job_description, resume_text):
    try:
        documents = [job_description, resume_text]
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf_matrix = vectorizer.fit_transform(documents)
        cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        score = cosine_sim * 100

        job_keywords = extract_keywords(job_description)
        resume_keywords = extract_keywords(resume_text)

        matched_keywords = {key: job_keywords[key] for key in resume_keywords if key in job_keywords}
        missing_keywords = {key: job_keywords[key] for key in job_keywords if key not in resume_keywords}

        return {
            'score': score,
            'found_keywords': matched_keywords,
            'missing_keywords': missing_keywords
        }
    except Exception as e:
        print(f"Error in calculate_score: {e}")
        return {'error': 'An error occurred while calculating the score.'}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'resume' not in request.files or 'job_description' not in request.form:
        return jsonify({'error': 'No file or job description provided'}), 400

    job_description = request.form['job_description']
    resume_file = request.files['resume']

    if resume_file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    filename = secure_filename(resume_file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    resume_file.save(filepath)

    if filename.endswith('.pdf'):
        resume_text = extract_text_from_pdf(filepath)
    elif filename.endswith('.docx'):
        resume_text = extract_text_from_docx(filepath)
    else:
        return jsonify({'error': 'Unsupported file type'}), 400

    result = calculate_score(job_description, resume_text)

    if 'error' in result:
        return jsonify(result), 500

    return jsonify({
        'score': result['score'],
        'found_keywords': dict(result['found_keywords']),
        'missing_keywords': dict(result['missing_keywords'])
    })

@app.route('/recalculate', methods=['POST'])
def recalculate_score():
    data = request.json
    job_description = data.get('job_description')
    found_keywords = data.get('found_keywords')

    if not job_description or not found_keywords:
        return jsonify({'error': 'Job description or found keywords missing'}), 400

    resume_text = ' '.join([f"{key} " * count for key, count in found_keywords.items()])

    try:
        result = calculate_score(job_description, resume_text)
        return jsonify({
            'score': result['score'],
            'found_keywords': dict(result['found_keywords']),
            'missing_keywords': dict(result['missing_keywords'])
        })
    except Exception as e:
        print(f"Error in recalculate_score: {e}")
        return jsonify({'error': 'An error occurred while recalculating the score.'}), 500

@app.route('/delete_keyword', methods=['POST'])
def delete_keyword():
    data = request.json
    job_description = data.get('job_description')
    found_keywords = data.get('found_keywords', {})
    missing_keywords = data.get('missing_keywords', {})
    keyword = data.get('keyword')
    table = data.get('table')

    if not job_description or not keyword or not table:
        return jsonify({'error': 'Missing data'}), 400

    if table == 'found':
        if keyword in found_keywords:
            del found_keywords[keyword]
    elif table == 'missing':
        if keyword in missing_keywords:
            del missing_keywords[keyword]

    try:
        resume_text = ' '.join([f"{key} " * count for key, count in found_keywords.items()])
        result = calculate_score(job_description, resume_text)
        return jsonify({
            'score': result['score'],
            'found_keywords': dict(result['found_keywords']),
            'missing_keywords': dict(result['missing_keywords'])
        })
    except Exception as e:
        print(f"Error in delete_keyword: {e}")
        return jsonify({'error': 'An error occurred while deleting the keyword.'}), 500

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    app.run(debug=True)
